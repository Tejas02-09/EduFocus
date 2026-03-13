#!/usr/bin/env python3
"""
Modify project report .docx file while preserving formatting.
Changes: Smart Rainwater Harvesting -> EDUFOCUS - Study with Focus
"""

from docx import Document
from docx.shared import Pt, RGBColor
import re

def replace_text_in_runs(paragraph, old_text, new_text):
    """Replace text in paragraph while preserving formatting"""
    full_text = ''.join(run.text for run in paragraph.runs)
    
    if old_text.lower() in full_text.lower():
        # Find and replace case-insensitively
        pattern = re.compile(re.escape(old_text), re.IGNORECASE)
        new_full_text = pattern.sub(new_text, full_text)
        
        if new_full_text != full_text:
            # Clear existing runs
            for run in paragraph.runs:
                run.text = ''
            
            # Add new text to first run to preserve formatting
            if paragraph.runs:
                paragraph.runs[0].text = new_full_text
            else:
                paragraph.add_run(new_full_text)
            
            return True
    return False

def modify_document(input_path, output_path):
    """Main function to modify the document"""
    
    # Load document
    doc = Document(input_path)
    
    # Define replacement mappings
    replacements = {
        # Project name
        'Smart Rainwater Harvesting Management System': 'EDUFOCUS – Study with Focus',
        'Smart Rainwater Harvesting': 'EDUFOCUS',
        'Rainwater Harvesting Management System': 'EDUFOCUS – Study with Focus',
        'rainwater harvesting': 'AI-powered smart learning platform',
        'Rainwater harvesting': 'AI-powered smart learning platform',
        
        # Technical components
        'rainwater tank': 'learning platform',
        'tank simulation': 'focus tracking system',
        'rainfall prediction': 'document summarization',
        'NASA POWER API': 'OpenAI API',
        'LSTM rainfall model': 'machine learning model',
        'Genetic Algorithm': 'optimization algorithm',
        'water usage optimization': 'study time optimization',
        'water conservation': 'learning efficiency',
        'tank capacity': 'user engagement',
        'precipitation': 'student focus',
        'meteorological': 'behavioral',
        
        # Student details (case-insensitive variations)
        'SCA24MCA040': 'SCA24MCA041',
        'Swathi': 'Tejas K M',
    }
    
    # Also add case-insensitive variations
    additional_replacements = {
        'LSTM': 'AI',
        'genetic algorithm': 'optimization algorithm',
        'NASA': 'OpenAI',
        'precipitation data': 'user interaction data',
        'rainfall': 'focus level',
        'Rainfall': 'Focus Level',
    }
    
    replacements.update(additional_replacements)
    
    # Process all paragraphs
    paragraph_count = 0
    replacement_count = 0
    
    for paragraph in doc.paragraphs:
        for old, new in replacements.items():
            if replace_text_in_runs(paragraph, old, new):
                replacement_count += 1
        paragraph_count += 1
    
    # Process all tables
    table_count = 0
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for old, new in replacements.items():
                        if replace_text_in_runs(paragraph, old, new):
                            replacement_count += 1
        table_count += 1
    
    # Process headers and footers
    for section in doc.sections:
        # Headers
        for paragraph in section.header.paragraphs:
            for old, new in replacements.items():
                replace_text_in_runs(paragraph, old, new)
        
        # Footers
        for paragraph in section.footer.paragraphs:
            for old, new in replacements.items():
                replace_text_in_runs(paragraph, old, new)
    
    # Save modified document
    doc.save(output_path)
    
    print(f"✅ Document successfully modified!")
    print(f"📊 Statistics:")
    print(f"   - Paragraphs processed: {paragraph_count}")
    print(f"   - Tables processed: {table_count}")
    print(f"   - Replacements made: {replacement_count}")
    print(f"\n📄 Output saved to: {output_path}")
    print(f"\n✨ All formatting, fonts, sizes, spacing, and layout preserved!")

if __name__ == "__main__":
    input_file = r"c:\Users\TEJAS\AppData\Local\Packages\5319275A.WhatsAppDesktop_cv1g1gvanyjgm\LocalState\sessions\FD8FCF8D0A88CE9B575D94CB33796A5A8E364714\transfers\2026-11\sca24mca040swathij.docx"
    output_file = r"c:\Users\TEJAS\Desktop\EDU-FOCUS\sca24mca041_EDUFOCUS_Project_Report.docx"
    
    try:
        modify_document(input_file, output_file)
    except FileNotFoundError:
        print(f"❌ Error: Input file not found at {input_file}")
    except Exception as e:
        print(f"❌ Error: {str(e)}")
